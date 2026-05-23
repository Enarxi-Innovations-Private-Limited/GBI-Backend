import os
import re
import zlib
import base64
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def render_mermaid_to_png(mermaid_code):
    """Compresses with zlib, encodes with urlsafe base64, and fetches a rendered PNG stream from Kroki."""
    try:
        compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f"https://kroki.io/mermaid/png/{encoded}"
        response = requests.get(url, timeout=25)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print(f"Kroki returned status {response.status_code}: {response.text}")
    except Exception as e:
        print(f"Diagram rendering error: {e}")
    return None

def add_formatted_runs(paragraph, text, default_color):
    """Splits markdown text into bold, italic, and inline code runs with Times New Roman typography."""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = default_color
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Times New Roman'
            run.font.italic = True
            run.font.color.rgb = default_color
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 50, 50)
        else:
            run = paragraph.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = default_color

def convert_md_to_docx(md_filepath, docx_filepath):
    print(f"\n--- Starting DOCX generation for {md_filepath} ---")
    if not os.path.exists(md_filepath):
        print(f"Error: File {md_filepath} not found.")
        return

    with open(md_filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 1. Page Margins (Standard 1 Inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 2. Executive Color Palette
    NAVY = RGBColor(27, 54, 93)      # Primary Headings
    SLATE = RGBColor(70, 80, 95)     # Secondary Headings / Rules
    CHARCOAL = RGBColor(40, 40, 40)  # Body Text

    # Normal Style Typography
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = CHARCOAL

    lines = content.split('\n')
    in_code_block = False
    code_buffer = []
    code_lang = ""

    for line in lines:
        raw_line = line.strip()

        # Handle Code and Mermaid Blocks
        if raw_line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = raw_line[3:].strip()
                code_buffer = []
            else:
                in_code_block = False
                block_text = "\n".join(code_buffer)
                
                if code_lang == 'mermaid' or code_lang == 'mermaid\n':
                    print("Rendering Mermaid diagram via Kroki API...")
                    img = render_mermaid_to_png(block_text)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    if img:
                        p.add_run().add_picture(img, width=Inches(6.0))
                        print(" -> Diagram successfully rendered & embedded!")
                    else:
                        r = p.add_run("[ Rendered Diagram Placeholder - Error Contacting Kroki API ]")
                        r.font.name = 'Times New Roman'
                        r.font.italic = True
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.right_indent = Inches(0.2)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    r = p.add_run(f"[{code_lang.upper() if code_lang else 'CODE'}]\n{block_text}")
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(30, 30, 30)
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # Handle Horizontal Rules
        if raw_line == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run("________________________________________________________________")
            r.font.name = 'Times New Roman'
            r.font.color.rgb = SLATE
            r.font.bold = True
            continue

        # Handle Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            
            p = doc.add_heading(level=min(level, 4))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_runs(p, title, NAVY if level <= 2 else SLATE)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.bold = True
            if level == 1:
                p.runs[0].font.size = Pt(22)
            elif level == 2:
                p.runs[0].font.size = Pt(16)
            else:
                p.runs[0].font.size = Pt(13)
            continue

        # Handle Bullet Lists
        if raw_line.startswith(('* ', '- ')):
            text = raw_line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, text, CHARCOAL)
            continue

        # Standard Paragraphs
        if raw_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_runs(p, line.strip(), CHARCOAL)

    try:
        doc.save(docx_filepath)
        print(f"Successfully generated professional document: {os.path.abspath(docx_filepath)}")
    except PermissionError:
        print(f"Warning: Could not save {docx_filepath} because it is currently open in Microsoft Word. Close the file to update.")

if __name__ == "__main__":
    convert_md_to_docx("GBI_Master_E2E_System_Specification.md", "GBI_Master_E2E_System_Specification.docx")
